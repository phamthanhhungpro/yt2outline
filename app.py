import yt_dlp
import openai
import requests
from fastapi import FastAPI, HTTPException, Request
from pydantic import BaseModel
import json
from functools import lru_cache
import hashlib
from docx import Document
from fastapi.responses import FileResponse
import re
from datetime import datetime
import os

def sanitize_filename(filename: str) -> str:
    # Remove invalid characters
    return re.sub(r'[<>:"/\\|?*#]', '', filename)

app = FastAPI()

# Set your OpenAI API key
openai.api_key = os.getenv("OPENAI_API_KEY")

class OutlineRequest(BaseModel):
    youtube_url: str
    name: str
    outline: dict

@lru_cache(maxsize=128)
def get_existing_captions_text(url):
    ydl_opts = {
        'skip_download': True,  # Do not download the video
        'writeautomaticsub': True,  # Extract automatic subtitles
        'outtmpl': '%(id)s',  # Use video ID as filename
        'nocheckcertificate': True,  # Use cookies for authentication
    }
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(url, download=False)
        subtitles = info.get("subtitles") or info.get("automatic_captions")
        if subtitles and "en" in subtitles:
            # Fetch caption text from the subtitle URL
            caption_url = subtitles["en"][-1]["url"]
            response = requests.get(caption_url)
            response.raise_for_status()
            return response.text, info["title"]  # Return the raw caption text and video title
    return None, None

@lru_cache(maxsize=128)
def download_audio(url):
    ydl_opts = {
        'format': 'bestaudio/best',  # Download the best available audio
        'outtmpl': '%(id)s.%(ext)s',  # Save file as VIDEO_ID.format
        'postprocessors': [],  # Skip postprocessing
        'nocheckcertificate': True,  # Use cookies for authentication
    }
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(url, download=True)
        return f"{info['id']}.{info['ext']}", info["title"]  # Return the downloaded file name and video title

@lru_cache(maxsize=128)
def generate_outline_cached(youtube_url: str):
    # Check for existing captions
    captions_text, video_title = get_existing_captions_text(youtube_url)
    if not captions_text:
        # Download audio file
        audio_file, video_title = download_audio(youtube_url)

        # Transcribe audio using OpenAI Whisper API
        with open(audio_file, "rb") as file:
            transcription_response = openai.audio.transcriptions.create(
                model="whisper-1",
                file=file,
                response_format="verbose_json",
                timestamp_granularities=["word"]
            )
        captions_text = transcription_response.text
        print(captions_text)

    def split_text(text, max_length):
        """Chia nội dung dài thành các phần nhỏ hơn"""
        return [text[i:i+max_length] for i in range(0, len(text), max_length)]

    # Chia transcription thành từng phần với độ dài giới hạn
    parts = split_text(captions_text, 1000)  # Giới hạn mỗi phần 4000 ký tự
    try:
    # Tạo dàn ý cho từng phần
        outlines = []
        for part in parts:
            prompt = f"""
            Keep original language. Create a most detail structured outline for the following video transcription. Format the response as JSON with the following structure:
            If the transcription does not have time ranges, you can help generate them based on the content. Be creative and informative.
            {{
                "sections": [
                    {{
                        "title": "Section Title",
                        "subsections": [
                            {{
                                "subtitle": "Subsection Title",
                                "timeRange": "Start Time - End Time"
                            }}
                        ]
                    }}
                ]
            }}

            Transcription:
            {part}
            """
            response = openai.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an assistant that generates outlines in JSON format for video content."},
                    {"role": "user", "content": prompt},
                ],
            )
            # Parse the JSON response for each part and append to outlines
            outline_part = json.loads(response.choices[0].message.content)
            outlines.extend(outline_part["sections"])  # Add sections directly to the combined list
    except Exception as e:
        print(e)
        raise HTTPException(status_code=500, detail="Error generating outline")

    return {
        "message": "Outline generated",
        "title": video_title,
        "outline": {
            "sections": outlines
        },
    }

@lru_cache(maxsize=128)
def generate_content_cached(youtube_url: str, outline_str: str):
    outline = json.loads(outline_str)
    sections = outline.get("sections", [])

    # Generate engaging content for each section
    content_list = []
    for section in sections:
        section_title = section["title"]
        prompt = f"""
        Create an engaging and captivating content for the following section of a video outline. Make sure the content is well-structured and interesting to read. Keep the language engaging and informative.
        The output language should be the same as the input language.

        Section Title: {section_title}
        Subsections: {json.dumps(section.get("subsections", []), indent=2)}
        """
        response = openai.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an assistant that generates engaging content for video outlines."},
                {"role": "user", "content": prompt},
            ],
        )
        content = response.choices[0].message.content
        content_list.append({
            "section_title": section_title,
            "content": content
        })

    return content_list

@app.get("/generate-outline/")
async def generate_outline(youtube_url: str):
    try:
        result = generate_outline_cached(youtube_url)
        return result
    except Exception as e:
        print(e)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate-docx/")
async def generate_docx(request: OutlineRequest, req: Request):
    try:
        outline_str = json.dumps(request.outline, sort_keys=True)
        content_list = generate_content_cached(request.youtube_url, outline_str)

        # Create a new Document
        doc = Document()
        doc.add_heading('Video Content', 0)

        for section in content_list:
            doc.add_heading(section["section_title"], level=1)
            doc.add_paragraph(section["content"])

        # Save the document
        current_time = datetime.now().strftime("%Y%m%d%H%M%S")
        file_path = sanitize_filename(request.name) + "_" + current_time + ".docx"
        doc.save("data/" + file_path)
        # Generate the downloadable link dynamically
        download_link = f"{req.url.scheme}://{req.client.host}:{req.url.port}/download/{file_path}"
        return {
            "url": request.youtube_url,
            "title": request.name,
            "filePath": download_link
        }
    except Exception as e:
        print(e)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/download/{file_name}")
async def download_file(file_name: str):
    file_path = f"./data/{file_name}"
    return FileResponse(file_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=file_name)

@app.post("/generate-new-outline/")
async def generate_new_outline(request: OutlineRequest):
    try:
        # Use the provided outline to generate a new outline
        prompt = f"""
        Based on the following outline, create most detailed structured outline. Format the response as JSON with the following structure:
        Be creative and informative. Keep the language engaging and informative. 
        If the outline does not have time ranges, you can help generate them based on the content.
        {{
            "sections": [
                {{
                    "title": "Section Title",
                    "subsections": [
                        {{
                            "subtitle": "Subsection Title",
                            "timeRange": "Start Time - End Time"
                        }}
                    ]
                }}
            ]
        }}

        Provided Outline:
        {json.dumps(request.outline, indent=2)}
        """
        response = openai.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an assistant that generates outlines in JSON format for video content."},
                {"role": "user", "content": prompt},
            ],
        )
        new_outline = json.loads(response.choices[0].message.content)

        return {
            "message": "New outline generated",
            "videoUrl": request.youtube_url,
            "title": request.name,
            "outline": {
                "sections": new_outline["sections"]
                }
        }
    except Exception as e:
        print(e)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate-seo-article/")
async def generate_seo_article(request: OutlineRequest, req: Request):
    try:
        # Use the provided outline to generate an SEO-optimized article
        prompt = f"""
        Based on the following outline, write a detailed SEO-optimized article. Make sure the content is well-structured, informative, and includes relevant keywords. The article should be ready to copy and paste for posting. Write in a natural, engaging style.
        Don't include time ranges in the article. Keep the language same as the input language.
        Outline:
        {json.dumps(request.outline, indent=2)}
        """
        response = openai.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an assistant that writes SEO-optimized articles based on provided outlines."},
                {"role": "user", "content": prompt},
            ],
        )
        article_content = response.choices[0].message.content.strip()

        # Create a new Document
        doc = Document()
        doc.add_heading(request.name, 0)
        doc.add_paragraph(article_content)

        # Save the document
        current_time = datetime.now().strftime("%Y%m%d%H%M%S")
        file_path = "data/" + sanitize_filename(request.name) + "_" + current_time + ".docx"
        doc.save(file_path)

        # Generate the downloadable link dynamically
        download_link = f"{req.url.scheme}://{req.client.host}:{req.url.port}/download/{file_path}"

        return {
            "message": "SEO-optimized article generated",
            "url": request.youtube_url,
            "title": request.name,
            "filePath": download_link
        }
    except Exception as e:
        print(e)
        raise HTTPException(status_code=500, detail=str(e))